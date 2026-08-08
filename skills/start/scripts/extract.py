#!/usr/bin/env python3
"""
Provenance-preserving extractor for a business document corpus.

It does not summarise and does not interpret: it normalises. It produces text
blocks, each labelled with its document and position (slide N, page N), because
without exact provenance the extraction is unusable: you have to go back to the
original slide every time a claim needs checking.

It also flags the pages that need visual inspection. On a sales deck this is
the part that matters most: the architectural promise is usually not written,
it is drawn. Three boxes with arrows saying "one single platform" produce no
extractable text at all and are a tenancy constraint.

Usage:
  python extract.py <file-o-cartella> [...] -o out/
  python extract.py corpus/ -o out/ --jsonl        one line per block
  python extract.py corpus/ -o out/ --min-chars 40 threshold for "text-poor page"

Output in out/:
  extract.md      blocks with a provenance heading, readable
  extract.jsonl   one JSON line per block (with --jsonl)
  inventory.json  inventory of documents and pages needing visual inspection
  render/         images of the flagged pages, ready to look at

Dependencies: check them before ingesting. With one missing, the corresponding
format produces zero blocks and a note that is easy not to read.

  pip install "markitdown[pptx,docx,pdf]" python-docx      pptx and docx
  poppler                                                  pdf
  LibreOffice (optional)                                   legacy .ppt/.doc,
                                                           rasterising pptx
"""

from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
from dataclasses import dataclass, asdict, field
from pathlib import Path

SUPPORTED = {".pptx", ".potx", ".ppt", ".pdf", ".docx", ".doc", ".md", ".txt"}
SLIDE_MARK = re.compile(r"<!--\s*Slide number:\s*(\d+)\s*-->", re.I)


@dataclass
class Block:
    source: str          # file name
    locator: str         # "slide 4", "page 12", "document"
    text: str
    kind: str = "text"   # text | notes | table


@dataclass
class DocInfo:
    source: str
    kind: str
    units: int = 0                 # slides or pages
    chars: int = 0
    has_text_layer: bool = True
    visual_review: list[int] = field(default_factory=list)
    rendered: int = 0              # how many flagged pages were rasterised
    note: str = ""


def run(cmd: list[str]) -> tuple[int, str]:
    # Explicit `encoding`, not the locale's: poppler and markitdown emit UTF-8, and on
    # Windows the locale is cp1252. With the implicit codec a page containing a curly
    # quote raises UnicodeDecodeError inside subprocess's reader thread, where the
    # exception does not propagate: the page comes back empty and gets flagged as
    # "text-poor", which is to say as graphical content. The tool then sends you to look
    # at an image that does not exist instead of telling you it could not read, and a
    # mutilated corpus passes for a complete one.
    try:
        p = subprocess.run(cmd, capture_output=True, text=True, timeout=300,
                           encoding="utf-8", errors="replace")
        return p.returncode, p.stdout
    except (subprocess.TimeoutExpired, FileNotFoundError) as e:
        return 1, f"{e}"


# ── PPTX ────────────────────────────────────────────────────────────────────

def extract_pptx(path: Path, out: Path, min_chars: int) -> tuple[list[Block], DocInfo]:
    src = path.name
    if path.suffix.lower() == ".ppt":                       # legacy: convert first
        rc, _ = run(["soffice", "--headless", "--convert-to", "pptx",
                     "--outdir", str(out / "_conv"), str(path)])
        cand = out / "_conv" / (path.stem + ".pptx")
        if rc == 0 and cand.exists():
            path = cand
        else:
            return [], DocInfo(src, "pptx", note="conversion from .ppt failed")

    rc, md = run(["markitdown", str(path)])
    if rc != 0 or not md.strip():
        return [], DocInfo(src, "pptx", note="markitdown produced no output")

    blocks, info = [], DocInfo(src, "pptx")
    parts = SLIDE_MARK.split(md)
    # parts = [preamble, n1, text1, n2, text2, ...]
    for i in range(1, len(parts) - 1, 2):
        n, body = int(parts[i]), parts[i + 1].strip()
        info.units = max(info.units, n)
        info.chars += len(body)
        if body:
            blocks.append(Block(src, f"slide {n}", body))
        if len(body) < min_chars:
            info.visual_review.append(n)     # text-poor slide: it is a picture
    if not blocks and md.strip():
        blocks.append(Block(src, "document", md.strip()))
    if info.visual_review:
        info.note = (f"{len(info.visual_review)} slides under {min_chars} characters: "
                     "likely graphical content. Worth looking at.")
    return blocks, info


# ── PDF ─────────────────────────────────────────────────────────────────────

def extract_pdf(path: Path, out: Path, min_chars: int) -> tuple[list[Block], DocInfo]:
    src = path.name
    info = DocInfo(src, "pdf")

    _, meta = run(["pdfinfo", str(path)])
    m = re.search(r"^Pages:\s+(\d+)", meta, re.M)
    info.units = int(m.group(1)) if m else 0

    _, fonts = run(["pdffonts", str(path)])
    info.has_text_layer = len(fonts.strip().splitlines()) > 2
    if not info.has_text_layer:
        info.note = "no text layer: scanned PDF. Look at it page by page, or run OCR."
        info.visual_review = list(range(1, min(info.units, 40) + 1))
        return [], info

    blocks = []
    for n in range(1, info.units + 1):
        rc, txt = run(["pdftotext", "-layout", "-f", str(n), "-l", str(n), str(path), "-"])
        body = (txt or "").strip()
        info.chars += len(body)
        if body:
            blocks.append(Block(src, f"page {n}", body))
        if len(body) < min_chars:
            info.visual_review.append(n)

    # A deck exported to PDF has little text on many pages: treat it as visual.
    if info.units and len(info.visual_review) > info.units * 0.4:
        info.note = ("many text-poor pages: likely a presentation exported to PDF. The "
                     "extracted text loses the layout, and on a sales deck the layout is "
                     "where the promise lives.")
    elif info.visual_review:
        info.note = f"{len(info.visual_review)} text-poor pages."
    return blocks, info


# ── DOCX ────────────────────────────────────────────────────────────────────

def extract_docx(path: Path, out: Path, min_chars: int) -> tuple[list[Block], DocInfo]:
    """python-docx rather than markitdown: markitdown has no [docx] extra here, and the
    native headings give better provenance than a split on a regex."""
    src = path.name
    if path.suffix.lower() == ".doc":
        rc, _ = run(["soffice", "--headless", "--convert-to", "docx",
                     "--outdir", str(out / "_conv"), str(path)])
        cand = out / "_conv" / (path.stem + ".docx")
        if rc == 0 and cand.exists():
            path = cand
        else:
            return [], DocInfo(src, "docx", note="conversion from .doc failed")

    try:
        import docx
    except ImportError:
        return [], DocInfo(src, "docx", note="needs python-docx: pip install python-docx")

    doc = docx.Document(str(path))
    info = DocInfo(src, "docx")
    blocks: list[Block] = []
    section, buf = "preamble", []

    def flush():
        body = "\n".join(buf).strip()
        if body:
            blocks.append(Block(src, f"§ {section}", body))
            info.chars += len(body)

    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        if (p.style.name or "").startswith("Heading"):
            flush()
            section, buf = txt, []
        else:
            buf.append(txt)
    flush()

    # The tables of a requirements analysis often hold the real requirements
    for n, tbl in enumerate(doc.tables, 1):
        rows = [" | ".join(c.text.strip() for c in r.cells) for r in tbl.rows]
        body = "\n".join(r for r in rows if r.strip(" |"))
        if body:
            blocks.append(Block(src, f"table {n}", body, kind="table"))
            info.chars += len(body)

    info.units = len(blocks)
    if not blocks:
        info.note = "no text extracted: empty document, or entirely images"
    return blocks, info


def extract_plain(path: Path, out: Path, min_chars: int) -> tuple[list[Block], DocInfo]:
    # `min_chars` is deliberately unused here, and the signature keeps it only so every
    # handler is called the same way. It marks a page or slide as text-poor, meaning the
    # content is probably in a diagram somebody has to look at. A plain file has no pages
    # and nothing to look at: a two line note is a short claim, not a suspect one, and
    # dropping it would lose the claim while the inventory reported the document as read.
    txt = path.read_text(encoding="utf-8", errors="replace")
    return ([Block(path.name, "document", txt.strip())] if txt.strip() else [],
            DocInfo(path.name, path.suffix.lstrip("."), units=1, chars=len(txt)))


HANDLERS = {
    ".pptx": extract_pptx, ".potx": extract_pptx, ".ppt": extract_pptx,
    ".pdf": extract_pdf,
    ".docx": extract_docx, ".doc": extract_docx,
    ".md": extract_plain, ".txt": extract_plain,
}


def build_extract_md(n_files: int, blocks: list[Block], silent: list[DocInfo],
                     infos: list[DocInfo]) -> str:
    """The report an agent reads. A pure function so it can be checked without poppler.

    Everything the extractor could not get is stated here, above everything it did get.
    `inventory.json` has always carried the same accounting, but nothing tells you to open
    it, and a gap you have to know to go looking for reads as completeness: a scanned PDF
    waiting for OCR and a file the filesystem refused both leave a corpus that looks whole.
    """
    md = ["# Business corpus extraction", "",
          "Generated by `extract.py`. This is not a framework artifact: it is raw material "
          "waiting to be classified. Where each claim belongs is decided with the routing "
          "table.", ""]

    if silent:
        md += [f"## {len(silent)} of {n_files} documents produced no text", "",
               "Nothing below came from these. Each one is a claim you do not have yet, "
               "not a claim that does not exist. Settle them before treating the "
               "extraction as the corpus.", "",
               "| Document | Why |", "|---|---|"]
        md += [f"| {i.source} | {i.note.replace('|', '/') or 'no text content'} |"
               for i in silent]
        md += [""]

    # One step further in: these produced text, so they are not silent, but the part that
    # matters is drawn rather than written. On a sales deck the architectural constraint is
    # almost always a diagram. This warning used to go to stdout and nowhere else, so it
    # survived exactly as long as the terminal did.
    need = [i for i in infos if i.visual_review]
    if need:
        md += [f"## {len(need)} documents have pages you have to look at", "",
               "These gave text, and the text is not the whole claim: a page this thin "
               "usually carries a diagram. Read the pages listed before classifying "
               "anything that came from these documents.", "",
               "| Document | Pages or slides |", "|---|---|"]
        md += [f"| {i.source} | {', '.join(str(p) for p in i.visual_review[:12])}"
               f"{' ...' if len(i.visual_review) > 12 else ''} |" for i in need]
        md += [""]

    for b in blocks:
        md += [f"## {b.source} — {b.locator}", "", b.text, ""]
    return "\n".join(md)


def render_pages(path: Path, pages: list[int], out: Path) -> list[str]:
    """Rasterise the flagged pages, so the agent can actually look at them."""
    if not pages or path.suffix.lower() != ".pdf":
        return []
    d = out / "render"
    d.mkdir(parents=True, exist_ok=True)
    made = []
    for n in pages[:40]:                       # a cap: looking at 200 pages is not a plan
        prefix = d / f"{path.stem}-p{n:03d}"
        rc, _ = run(["pdftoppm", "-jpeg", "-r", "150", "-f", str(n), "-l", str(n),
                     str(path), str(prefix)])
        if rc == 0:
            made += [str(p) for p in sorted(d.glob(f"{path.stem}-p{n:03d}*.jpg"))]
    return made


def main() -> int:
    # Same trap as the validator: the Windows console is not UTF-8, and a character
    # outside the codepage ends the extraction with a traceback after the files have
    # already been written. The result exists but you never get to read it.
    for stream in (sys.stdout, sys.stderr):
        if hasattr(stream, "reconfigure"):
            stream.reconfigure(errors="replace")

    ap = argparse.ArgumentParser(description="Extract a business corpus, preserving provenance")
    ap.add_argument("inputs", nargs="+", type=Path)
    ap.add_argument("-o", "--out", type=Path, default=Path("ingest-out"),
                    help="output DIRECTORY, created if absent")
    ap.add_argument("--jsonl", action="store_true")
    ap.add_argument("--min-chars", type=int, default=40,
                    help="a page or slide under this many characters is text-poor and gets "
                         "flagged for you to look at. Pages and slides only: a short .md or "
                         ".txt is a short claim, not a suspect one, and is kept whole")
    ap.add_argument("--no-render", action="store_true")
    args = ap.parse_args()

    files: list[Path] = []
    for i in args.inputs:
        if i.is_dir():
            # The README of a corpus/ folder explains what to put in it: the framework
            # wrote it, not the business. Ingesting it would put our own instructions
            # into ING.md wearing the shape of a customer claim.
            files += [p for p in sorted(i.rglob("*"))
                      if p.suffix.lower() in SUPPORTED and p.name.lower() != "readme.md"]
        elif i.suffix.lower() in SUPPORTED:
            files.append(i)
        else:
            print(f"! skipped (unhandled format): {i}", file=sys.stderr)
    if not files:
        print("No handleable document found.", file=sys.stderr)
        return 1

    # `-o` names a directory. Pointed at a file it used to end in a FileExistsError
    # traceback, which reads as a crash in the extractor rather than as a typo in the
    # command, and the difference matters when an agent has to decide whether to retry.
    if args.out.exists() and not args.out.is_dir():
        print(f"-o wants a directory and {args.out} is a file.", file=sys.stderr)
        return 1
    args.out.mkdir(parents=True, exist_ok=True)
    all_blocks: list[Block] = []
    infos: list[DocInfo] = []
    silent: list[DocInfo] = []

    for f in files:
        h = HANDLERS[f.suffix.lower()]
        try:
            blocks, info = h(f, args.out, args.min_chars)
        except Exception as e:                        # one broken document must not stop the batch
            blocks, info = [], DocInfo(f.name, f.suffix.lstrip("."), note=f"error: {e}")
        if not args.no_render and info.visual_review:
            rendered = render_pages(f, info.visual_review, args.out)
            info.rendered = len(rendered)
            if rendered:
                info.note += f" Images in {args.out / 'render'}/ ({len(rendered)})."
        all_blocks += blocks
        infos.append(info)
        if not blocks:
            silent.append(info)
        print(f"- {f.name}: {len(blocks)} blocks, {info.units} units"
              + (f" - {info.note}" if info.note else ""))

    (args.out / "extract.md").write_text(
        build_extract_md(len(files), all_blocks, silent, infos), encoding="utf-8")

    if args.jsonl:
        with (args.out / "extract.jsonl").open("w", encoding="utf-8") as fh:
            for b in all_blocks:
                fh.write(json.dumps(asdict(b), ensure_ascii=False) + "\n")

    (args.out / "inventory.json").write_text(
        json.dumps({"documents": [asdict(i) for i in infos],
                    "blocks": len(all_blocks)}, indent=2, ensure_ascii=False),
        encoding="utf-8")

    print(f"\n{len(all_blocks)} blocks from {len(files)} documents -> {args.out}/extract.md")
    if silent:
        print(f"{len(silent)} of them gave no text, listed at the top of extract.md")
    need = [i for i in infos if i.visual_review]
    if need:
        print(f"\n{len(need)} documents need visual inspection:")
        for i in need:
            print(f"  - {i.source}: pages/slides {i.visual_review[:12]}"
                  f"{' ...' if len(i.visual_review) > 12 else ''}")
        print("\nOn a sales deck the architectural constraint is usually drawn rather "
              "than written: look at these pages before classifying anything.")

        # Pointing at a folder that was never produced teaches people to ignore the
        # warning, and this is the one part of the extraction you cannot delegate.
        if any(i.rendered for i in need):
            print(f"  Images ready in {args.out / 'render'}/")
        manual = [i for i in need if not i.rendered]
        if manual:
            print("  Not rasterisable here - open them by hand at the pages listed:")
            for i in manual:
                print(f"    - {i.source}")
            if any(i.kind in {"pptx", "docx"} for i in manual):
                print("  (rasterising pptx and docx too needs LibreOffice: "
                      "https://www.libreoffice.org/download)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
