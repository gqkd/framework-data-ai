#!/usr/bin/env python3
"""
Estrattore con provenienza per il corpus di documenti business.

Non riassume e non interpreta: normalizza. Produce blocchi di testo ciascuno
etichettato con documento e posizione (slide N, pagina N), perché senza la
provenienza esatta l'estrazione è inutilizzabile — dovrai tornare alla slide
originale ogni volta che un'affermazione va verificata.

Segnala inoltre le pagine che richiedono ispezione visiva. È la parte che conta
più di tutte su un deck commerciale: la promessa architetturale spesso non è
scritta, è disegnata. Tre box con delle frecce che dicono "piattaforma unica"
non producono nessun testo estraibile e sono un vincolo di tenancy.

Uso:
  python extract.py <file-o-cartella> [...] -o out/
  python extract.py corpus/ -o out/ --jsonl        una riga per blocco
  python extract.py corpus/ -o out/ --min-chars 40 soglia per "pagina povera di testo"

Output in out/:
  extract.md      blocchi con intestazione di provenienza, leggibile
  extract.jsonl   una riga JSON per blocco (con --jsonl)
  inventory.json  inventario dei documenti e pagine da ispezionare visivamente
  render/         immagini delle pagine segnalate, pronte da guardare

Dipendenze — verificale prima di ingestare, perché mancandone una il formato
corrispondente produce zero blocchi e una nota facile da non leggere:

  pip install "markitdown[pptx,docx,pdf]" python-docx      pptx e docx
  poppler                                                  pdf
  LibreOffice (facoltativo)                                .ppt/.doc legacy,
                                                           rasterizzazione di pptx
"""

from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
from dataclasses import dataclass, asdict, field
from pathlib import Path

SUPPORTED = {".pptx", ".potx", ".ppt", ".pdf", ".docx", ".doc", ".md", ".txt"}
SLIDE_MARK = re.compile(r"<!--\s*Slide number:\s*(\d+)\s*-->", re.I)


@dataclass
class Block:
    source: str          # nome del file
    locator: str         # "slide 4", "pagina 12", "documento"
    text: str
    kind: str = "text"   # text | notes | table


@dataclass
class DocInfo:
    source: str
    kind: str
    units: int = 0                 # slide o pagine
    chars: int = 0
    has_text_layer: bool = True
    visual_review: list[int] = field(default_factory=list)
    rendered: int = 0              # quante delle pagine segnalate sono state rasterizzate
    note: str = ""


def run(cmd: list[str]) -> tuple[int, str]:
    # `encoding` esplicito, non quello del locale: poppler e markitdown emettono
    # UTF-8, e su Windows il locale è cp1252. Con la codifica implicita una pagina
    # che contiene una virgoletta tipografica solleva UnicodeDecodeError dentro il
    # thread lettore di subprocess — dove l'eccezione non propaga: la pagina torna
    # vuota e viene segnalata come «povera di testo», cioè come contenuto grafico.
    # Lo strumento ti manda a guardare un'immagine che non esiste invece di dirti
    # che non ha saputo leggere, e il corpus mutilato passa per completo.
    try:
        p = subprocess.run(cmd, capture_output=True, text=True, timeout=300,
                           encoding="utf-8", errors="replace")
        return p.returncode, p.stdout
    except (subprocess.TimeoutExpired, FileNotFoundError) as e:
        return 1, f"{e}"


# ── PPTX ────────────────────────────────────────────────────────────────────

def extract_pptx(path: Path, out: Path, min_chars: int) -> tuple[list[Block], DocInfo]:
    src = path.name
    if path.suffix.lower() == ".ppt":                       # legacy: converti
        rc, _ = run(["soffice", "--headless", "--convert-to", "pptx",
                     "--outdir", str(out / "_conv"), str(path)])
        cand = out / "_conv" / (path.stem + ".pptx")
        if rc == 0 and cand.exists():
            path = cand
        else:
            return [], DocInfo(src, "pptx", note="conversione da .ppt fallita")

    rc, md = run(["markitdown", str(path)])
    if rc != 0 or not md.strip():
        return [], DocInfo(src, "pptx", note="markitdown non ha prodotto output")

    blocks, info = [], DocInfo(src, "pptx")
    parts = SLIDE_MARK.split(md)
    # parts = [preambolo, n1, testo1, n2, testo2, ...]
    for i in range(1, len(parts) - 1, 2):
        n, body = int(parts[i]), parts[i + 1].strip()
        info.units = max(info.units, n)
        info.chars += len(body)
        if body:
            blocks.append(Block(src, f"slide {n}", body))
        if len(body) < min_chars:
            info.visual_review.append(n)     # slide povera di testo → è un'immagine
    if not blocks and md.strip():
        blocks.append(Block(src, "documento", md.strip()))
    if info.visual_review:
        info.note = (f"{len(info.visual_review)} slide con meno di {min_chars} caratteri: "
                     "probabile contenuto grafico. Da guardare.")
    return blocks, info


# ── PDF ─────────────────────────────────────────────────────────────────────

def extract_pdf(path: Path, out: Path, min_chars: int) -> tuple[list[Block], DocInfo]:
    src = path.name
    info = DocInfo(src, "pdf")

    _, meta = run(["pdfinfo", str(path)])
    m = re.search(r"^Pages:\s+(\d+)", meta, re.M)
    info.units = int(m.group(1)) if m else 0

    _, fonts = run(["pdffonts", str(path)])
    info.has_text_layer = len(fonts.strip().splitlines()) > 2
    if not info.has_text_layer:
        info.note = "nessun layer di testo: PDF scansionato. Va guardato pagina per pagina o passato per OCR."
        info.visual_review = list(range(1, min(info.units, 40) + 1))
        return [], info

    blocks = []
    for n in range(1, info.units + 1):
        rc, txt = run(["pdftotext", "-layout", "-f", str(n), "-l", str(n), str(path), "-"])
        body = (txt or "").strip()
        info.chars += len(body)
        if body:
            blocks.append(Block(src, f"pagina {n}", body))
        if len(body) < min_chars:
            info.visual_review.append(n)

    # Un deck esportato in PDF ha poco testo su molte pagine: trattalo come visivo.
    if info.units and len(info.visual_review) > info.units * 0.4:
        info.note = ("molte pagine povere di testo: probabile presentazione esportata in PDF. "
                     "Il testo estratto perde il layout, e su un deck commerciale il layout "
                     "è dove sta la promessa.")
    elif info.visual_review:
        info.note = f"{len(info.visual_review)} pagine povere di testo."
    return blocks, info


# ── DOCX ────────────────────────────────────────────────────────────────────

def extract_docx(path: Path, out: Path, min_chars: int) -> tuple[list[Block], DocInfo]:
    """python-docx invece di markitdown: qui markitdown non ha l'extra [docx], e le
    intestazioni native danno una provenienza migliore di uno split su regex."""
    src = path.name
    if path.suffix.lower() == ".doc":
        rc, _ = run(["soffice", "--headless", "--convert-to", "docx",
                     "--outdir", str(out / "_conv"), str(path)])
        cand = out / "_conv" / (path.stem + ".docx")
        if rc == 0 and cand.exists():
            path = cand
        else:
            return [], DocInfo(src, "docx", note="conversione da .doc fallita")

    try:
        import docx
    except ImportError:
        return [], DocInfo(src, "docx", note="serve python-docx: pip install python-docx")

    doc = docx.Document(str(path))
    info = DocInfo(src, "docx")
    blocks: list[Block] = []
    section, buf = "premessa", []

    def flush():
        body = "\n".join(buf).strip()
        if body:
            blocks.append(Block(src, f"§ {section}", body))
            info.chars += len(body)

    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        if (p.style.name or "").startswith("Heading"):
            flush()
            section, buf = txt, []
        else:
            buf.append(txt)
    flush()

    # Le tabelle di un'analisi dei requisiti contengono spesso i requisiti veri
    for n, tbl in enumerate(doc.tables, 1):
        rows = [" | ".join(c.text.strip() for c in r.cells) for r in tbl.rows]
        body = "\n".join(r for r in rows if r.strip(" |"))
        if body:
            blocks.append(Block(src, f"tabella {n}", body, kind="table"))
            info.chars += len(body)

    info.units = len(blocks)
    if not blocks:
        info.note = "nessun testo estratto: documento vuoto o interamente in immagini"
    return blocks, info


def extract_plain(path: Path, out: Path, min_chars: int) -> tuple[list[Block], DocInfo]:
    txt = path.read_text(encoding="utf-8", errors="replace")
    return ([Block(path.name, "documento", txt.strip())] if txt.strip() else [],
            DocInfo(path.name, path.suffix.lstrip("."), units=1, chars=len(txt)))


HANDLERS = {
    ".pptx": extract_pptx, ".potx": extract_pptx, ".ppt": extract_pptx,
    ".pdf": extract_pdf,
    ".docx": extract_docx, ".doc": extract_docx,
    ".md": extract_plain, ".txt": extract_plain,
}


def render_pages(path: Path, pages: list[int], out: Path) -> list[str]:
    """Rasterizza le pagine segnalate, così l'agente può guardarle."""
    if not pages or path.suffix.lower() != ".pdf":
        return []
    d = out / "render"
    d.mkdir(parents=True, exist_ok=True)
    made = []
    for n in pages[:40]:                       # tetto: guardare 200 pagine non è una strategia
        prefix = d / f"{path.stem}-p{n:03d}"
        rc, _ = run(["pdftoppm", "-jpeg", "-r", "150", "-f", str(n), "-l", str(n),
                     str(path), str(prefix)])
        if rc == 0:
            made += [str(p) for p in sorted(d.glob(f"{path.stem}-p{n:03d}*.jpg"))]
    return made


def main() -> int:
    # Stessa insidia del validatore: la console di Windows non è UTF-8, e un
    # carattere fuori dalla codepage fa terminare l'estrazione con un traceback
    # dopo aver già scritto i file — l'esito c'è ma non lo leggi.
    for stream in (sys.stdout, sys.stderr):
        if hasattr(stream, "reconfigure"):
            stream.reconfigure(errors="replace")

    ap = argparse.ArgumentParser(description="Estrae il corpus business con provenienza")
    ap.add_argument("inputs", nargs="+", type=Path)
    ap.add_argument("-o", "--out", type=Path, default=Path("ingest-out"))
    ap.add_argument("--jsonl", action="store_true")
    ap.add_argument("--min-chars", type=int, default=40)
    ap.add_argument("--no-render", action="store_true")
    args = ap.parse_args()

    files: list[Path] = []
    for i in args.inputs:
        if i.is_dir():
            # Il README di una cartella corpus/ spiega cosa metterci dentro: lo ha
            # scritto il framework, non il business. Ingestarlo metterebbe le nostre
            # istruzioni in ING.md con la forma di un'affermazione del cliente.
            files += [p for p in sorted(i.rglob("*"))
                      if p.suffix.lower() in SUPPORTED and p.name.lower() != "readme.md"]
        elif i.suffix.lower() in SUPPORTED:
            files.append(i)
        else:
            print(f"! ignorato (formato non gestito): {i}", file=sys.stderr)
    if not files:
        print("Nessun documento gestibile trovato.", file=sys.stderr)
        return 1

    args.out.mkdir(parents=True, exist_ok=True)
    all_blocks: list[Block] = []
    infos: list[DocInfo] = []

    for f in files:
        h = HANDLERS[f.suffix.lower()]
        try:
            blocks, info = h(f, args.out, args.min_chars)
        except Exception as e:                        # un documento rotto non ferma il lotto
            blocks, info = [], DocInfo(f.name, f.suffix.lstrip("."), note=f"errore: {e}")
        if not args.no_render and info.visual_review:
            rendered = render_pages(f, info.visual_review, args.out)
            info.rendered = len(rendered)
            if rendered:
                info.note += f" Immagini in {args.out / 'render'}/ ({len(rendered)})."
        all_blocks += blocks
        infos.append(info)
        print(f"· {f.name}: {len(blocks)} blocchi, {info.units} unità"
              + (f" — {info.note}" if info.note else ""))

    md = ["# Estrazione del corpus business", "",
          "Generato da `extract.py`. Non è un artefatto del framework: è materiale grezzo "
          "da classificare. La destinazione di ogni affermazione si decide con "
          "`references/routing-table.md`.", ""]
    for b in all_blocks:
        md += [f"## {b.source} — {b.locator}", "", b.text, ""]
    (args.out / "extract.md").write_text("\n".join(md), encoding="utf-8")

    if args.jsonl:
        with (args.out / "extract.jsonl").open("w", encoding="utf-8") as fh:
            for b in all_blocks:
                fh.write(json.dumps(asdict(b), ensure_ascii=False) + "\n")

    (args.out / "inventory.json").write_text(
        json.dumps({"documents": [asdict(i) for i in infos],
                    "blocks": len(all_blocks)}, indent=2, ensure_ascii=False),
        encoding="utf-8")

    need = [i for i in infos if i.visual_review]
    print(f"\n{len(all_blocks)} blocchi da {len(files)} documenti -> {args.out}/extract.md")
    if need:
        print(f"\n{len(need)} documenti richiedono ispezione visiva:")
        for i in need:
            print(f"  · {i.source}: pagine/slide {i.visual_review[:12]}"
                  f"{' …' if len(i.visual_review) > 12 else ''}")
        print("\nSu un deck commerciale il vincolo architetturale è spesso disegnato, non "
              "scritto: queste pagine vanno guardate prima di classificare qualsiasi cosa.")

        # Rimandare a una cartella che non è stata prodotta insegna a ignorare la
        # segnalazione, che è l'unica parte dell'estrazione che non puoi delegare.
        if any(i.rendered for i in need):
            print(f"  Immagini pronte in {args.out / 'render'}/")
        manual = [i for i in need if not i.rendered]
        if manual:
            print("  Non rasterizzabili qui — aprili a mano alle pagine indicate:")
            for i in manual:
                print(f"    · {i.source}")
            if any(i.kind in {"pptx", "docx"} for i in manual):
                print("  (per rasterizzare anche pptx e docx serve LibreOffice: "
                      "https://www.libreoffice.org/download)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
